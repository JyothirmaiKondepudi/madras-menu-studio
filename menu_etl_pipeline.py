#!/usr/bin/env python3
"""
Menu ETL Pipeline — turns a folder of inconsistent Word/PDF menu documents
into a structured, deduplicated dish database.

This is NOT a RAG index, and it does NOT need dbt. It's a one-time
(re-runnable) batch ELT: raw docs -> per-doc structured JSON (via Claude)
-> merged/deduped/categorized seed data -> loaded straight into Postgres.
Categorization (cuisine, occasion, price tier, tax category, etc.) happens
in plain Python during the aggregate step below — no separate SQL
transformation layer needed at this data volume.

Six stages, run separately so a crash/rerun never re-pays for docs already done:

    python menu_etl_pipeline.py curate    --input ./menus-source/raw --output ./menus-source/curated
    python menu_etl_pipeline.py extract   --input ./menus-source/curated --output ./extracted
    python menu_etl_pipeline.py aggregate --input ./extracted --output ./seed
    python menu_etl_pipeline.py load      --input ./seed --database-url $DATABASE_URL
    python menu_etl_pipeline.py stage     --input ./seed --database-url $DATABASE_URL
    python menu_etl_pipeline.py promote   --database-url $DATABASE_URL

`load` and `stage`/`promote` are two alternative review paths for the same
aggregate output, not a required sequence — use whichever review method fits:
`load` goes straight from a human-reviewed `dish_review.csv` (opened in
Excel) into the live `menu_items` table. `stage` instead pushes aggregate's
output into `menu_item_drafts` (review_status='pending') for review through
the app's own /review web page (card grid, approve/reject/edit) instead of a
spreadsheet; `promote` then upserts whatever got approved/edited there into
the same live `menu_items` table `load` targets directly. Nothing not
explicitly reviewed in either path ever reaches menu_items.

`curate` exists because a real source archive (e.g. a Google Drive export
going back over a decade) is not just menus — it's mixed in with plain junk
(0-byte files, Word `~$` lock files), exact-duplicate copies of the same
doc saved under different names/folders, and documents that are real files
but never contain a menu at all (insurance certificates, W-9s, name-tag
templates, floor plans, rental-equipment lists). Every one of those still
costs a full `extract` API call if not filtered out first, for nothing.
Deliberately NOT filtered by filename category (e.g. "invoice"-sounding
names): real per-event menus have turned up under invoice/pickup/plain
client-name filenames in this archive, so category guesses are excluded
only when the category itself could never contain food content — see
CURATE_EXCLUDE_PATTERNS below. Everything else, however unpromising the
name looks, goes on to `extract` and lets the actual extraction call decide
(cheaply — a doc with no real menu just comes back with an empty `menus`
list).

Requires:
    pip install anthropic python-docx pdfplumber rapidfuzz psycopg2-binary python-dotenv
    ANTHROPIC_API_KEY and DATABASE_URL (only needed for the load step) read
    automatically from this project's .env file — no manual export needed.

The `load` step writes into tables matching prisma/schema.prisma (included
alongside this script) — run `prisma migrate dev` on that schema first so
the tables exist, then run `load` to populate them.

Why an LLM and not regex: your source docs have no consistent structure —
bold headers, asterisk dividers, multiple unrelated menus concatenated in
one file, no explicit veg/non-veg markers, and pricing that's per-menu-
package rather than per-dish. A single doc can contain three full menus
(see "Kerala Inspired Menus.docx": three distinct wedding cocktail/dinner
menus in one file, priced at $47/$68/$110 per person). Regex can't reliably
tell where one menu ends and the next begins, or that "Chicken 65" is
non-veg and South Indian while "Palak Methi Corn Paneer" is veg and North
Indian. An LLM extraction pass, constrained to a strict JSON schema and a
controlled cuisine-tag vocabulary, handles that variance directly.
"""

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path

from dotenv import load_dotenv

# Python doesn't read .env files on its own the way Next.js's `dotenv/config`
# import does — load the same .env this project's Node side already uses, so
# ANTHROPIC_API_KEY/DATABASE_URL only need to live in one place.
load_dotenv()

# ---------------------------------------------------------------------------
# Controlled vocabulary — keep this in sync with your app's schema so tags
# extracted here drop straight into the real database later.
# ---------------------------------------------------------------------------

CUISINE_TAGS = [
    "north_indian", "south_indian", "punjabi", "sindhi", "marathi",
    "pakistani", "bangladeshi", "gujarati", "kerala", "telugu_andhra",
    "tamil", "rajasthani", "fusion", "chinese", "mexican", "italian",
    "american", "mediterranean", "thai", "sushi", "modern",
]

COURSES = [
    "appetizer", "live_station", "main", "bread", "rice_biryani",
    "side", "salad", "dessert", "beverage", "condiment",
]

STAPLE_HINTS = {
    "chai", "coffee", "naan", "rice", "dal", "papad", "raita", "pickle",
    "water", "lime", "onion", "cilantro",
}

ALLERGENS = ["nuts", "dairy", "gluten", "shellfish", "egg", "soy", "sesame"]

DIETARY_FLAGS = ["vegan", "jain", "halal", "kosher", "gluten_free"]

RELIGION_SUITABILITY = ["hindu", "muslim", "christian", "any"]

OCCASION_SUITABILITY = [
    "breakfast", "lunch", "wedding_lunch", "welcome_dinner", "mehndi",
    "sangeet", "cocktail_hour", "dinner_reception", "late_night_snacks",
    "dessert_only", "any",
]

TAX_CATEGORIES = ["prepared_food", "alcohol", "equipment_rental", "service_labor"]

SPICE_LEVELS = ["mild", "medium", "hot"]

EXTRACTION_SCHEMA = {
    "name": "record_menus",
    "description": "Structured extraction of every distinct menu and dish found in a catering document.",
    "input_schema": {
        "type": "object",
        "properties": {
            "menus": {
                "type": "array",
                "description": "One entry per distinct menu/package found in the document. A single doc often contains several (e.g. 'Menu 1', 'Menu 2', 'Menu 3').",
                "items": {
                    "type": "object",
                    "properties": {
                        "menu_label": {"type": "string", "description": "The menu's own heading, verbatim, e.g. 'Wedding Cocktail and Dinner Reception Menu 2'."},
                        "occasion_type_guess": {"type": "string", "description": "Best guess: welcome_dinner, breakfast, mehndi, sangeet, wedding_lunch, cocktail_hour, dinner_reception, late_night_snacks, or other."},
                        "price_per_person": {"type": ["number", "null"], "description": "Per-person price if stated, else null."},
                        "other_charges": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "label": {"type": "string"},
                                    "amount": {"type": "number"},
                                },
                                "required": ["label", "amount"],
                            },
                            "description": "Non-food line items: transportation, staffing, equipment, etc.",
                        },
                        "items": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "name": {"type": "string", "description": "Dish name, cleaned of stray formatting/asterisks."},
                                    "course": {"type": "string", "enum": COURSES},
                                    "veg_nonveg_guess": {"type": "string", "enum": ["veg", "nonveg", "unclear"]},
                                    "cuisine_tags_guess": {"type": "array", "items": {"type": "string", "enum": CUISINE_TAGS}},
                                    "price_weight_guess": {"type": "string", "enum": ["light", "standard", "premium"], "description": "Rough ingredient-cost tier — premium for lamb/shrimp/specialty items, light for a plain side or condiment."},
                                    "is_staple_guess": {"type": "boolean", "description": "True for items like chai, coffee, naan, plain rice/dal that can repeat across every menu without counting against variety."},
                                    "allergens_guess": {"type": "array", "items": {"type": "string", "enum": ALLERGENS}},
                                    "dietary_flags_guess": {"type": "array", "items": {"type": "string", "enum": DIETARY_FLAGS}, "description": "Only include a flag if the dish plausibly qualifies as-is (e.g. a plain dal is vegan; don't guess halal/kosher unless the meat prep implies it)."},
                                    "religion_suitability_guess": {"type": "array", "items": {"type": "string", "enum": RELIGION_SUITABILITY}, "description": "'any' unless the dish is clearly unsuitable for a tradition (e.g. beef dishes excluded from Hindu weddings, pork excluded from Muslim events)."},
                                    "occasion_suitability_guess": {"type": "array", "items": {"type": "string", "enum": OCCASION_SUITABILITY}, "description": "Which occasion types this dish fits; 'any' if broadly appropriate."},
                                    "spice_level_guess": {"type": "string", "enum": SPICE_LEVELS},
                                    "prep_method_guess": {"type": "string", "description": "e.g. fried, grilled, steamed, tandoor, live-station-flambe."},
                                    "tax_category_guess": {"type": "string", "enum": TAX_CATEGORIES},
                                    "confidence": {"type": "string", "enum": ["high", "medium", "low"], "description": "Your confidence in the veg/nonveg and cuisine guesses for this specific item."},
                                },
                                "required": ["name", "course", "veg_nonveg_guess", "cuisine_tags_guess", "price_weight_guess", "is_staple_guess", "allergens_guess", "dietary_flags_guess", "religion_suitability_guess", "occasion_suitability_guess", "tax_category_guess", "confidence"],
                            },
                        },
                    },
                    "required": ["menu_label", "occasion_type_guess", "items"],
                },
            }
        },
        "required": ["menus"],
    },
}

EXTRACTION_PROMPT = """You are extracting structured dish data from a catering company's internal menu document. The formatting is inconsistent — bold headers, asterisk dividers, multiple unrelated menus concatenated in one file, prices that apply to a whole menu package rather than per dish, and no explicit veg/non-veg labels (infer from the dish name and known ingredients).

Find every distinct menu in the document (a doc frequently contains 2-3 separate menus, e.g. "Menu 1" / "Menu 2" / "Menu 3" at different price points) and every dish within each menu. Skip pure boilerplate (phone numbers, page numbers, the company footer).

A very common mistake to avoid: station/display/condiment-bar sections list several standalone, separately-named foods on one line — joined by a comma, "&", "/", "and", or a dash — and it is NEVER one combined item, no matter which food category it is (rice, noodles, beverages, sauces, garnishes, chutneys, anything). Each named food in that list is its own separate item. This rule applies uniformly across every course/category — don't treat it as something that only applies to the kind of food in whatever example follows.

Examples across different categories, all the same underlying mistake:
- "Lemon Rice, Jeera Rice, Tomato Rice, Curd Rice Displayed Traditional Mud Pots" → FOUR separate rice dishes
- "Hakka Noodles & Vegetable Fried Rice" → TWO separate dishes
- "Freshly Brewed Masala Chai/Madras Coffee" → TWO separate beverages
- "Condiments station with Chopped Onion, Fresh cut Cilantro, Tomato, Yogurt, Punjabi Pickle, Green Chilies, and Fresh Cut Lime Wedges" → SEVEN separate garnishes/condiments
- "Guest's Choice of Sauces - Teriyaki, Stir Fry, Chilli Garlic Sauce, Lime Ginger Sauce" → FOUR separate sauces

Only keep foods together as a single item when they genuinely form one composed dish a guest would order as a unit (e.g. "Chicken Tikka Masala with Basmati Rice" as a single plated entrée) — not when a line is simply listing multiple standalone options side by side, regardless of what category those options belong to.

Use the record_menus tool to return your extraction.

Document filename: {filename}

Document text:
---
{doc_text}
---"""


def extract_text_from_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def extract_text_from_pdf(path: Path) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    raise ValueError(f"Unsupported file type: {path}")


def call_claude_extract(client, filename: str, doc_text: str) -> dict:
    # Trim absurdly long docs defensively; most menu docs are a few pages.
    doc_text = doc_text[:40000]

    # Identity-linked/service-account API keys can reach more than one
    # workspace, so the API needs to be told which workspace each request
    # acts in via this header. A plain key created from inside one specific
    # workspace doesn't need it — so only send it when it's actually set.
    extra_headers = {}
    workspace_id = os.environ.get("ANTHROPIC_WORKSPACE_ID")
    if workspace_id:
        extra_headers["anthropic-workspace-id"] = workspace_id

    # A doc that packs several full menus (common for a "complete celebration"
    # file with sangeet + lunch + reception all in one) needs well over the
    # old 8000-token cap to describe every item under our (verbose,
    # 12-field-per-item) schema — 8000 truncated mid-JSON with no error and
    # silently wrote back an empty extraction. Even 32000 wasn't enough for
    # the largest observed doc (a 4-menu, 100+ item file), so 64000; .stream()
    # is required by the SDK once max_tokens is high enough that a
    # non-streaming call could exceed its timeout.
    MAX_TOKENS = 64000
    with client.messages.stream(
        model="claude-sonnet-4-5",
        max_tokens=MAX_TOKENS,
        tools=[EXTRACTION_SCHEMA],
        tool_choice={"type": "tool", "name": "record_menus"},
        messages=[{
            "role": "user",
            "content": EXTRACTION_PROMPT.format(filename=filename, doc_text=doc_text),
        }],
        extra_headers=extra_headers,
    ) as stream:
        resp = stream.get_final_message()

    if resp.stop_reason == "max_tokens":
        # Don't silently write back a truncated/partial extraction — surface
        # it so cmd_extract's per-doc try/except logs a visible FAILED line
        # instead of a quietly empty result.
        raise RuntimeError(
            f"Extraction for {filename} hit the {MAX_TOKENS}-token cap before finishing "
            f"(doc likely has an unusually large number of menus/items) — "
            f"needs a bigger max_tokens or splitting the doc before extraction."
        )
    for block in resp.content:
        if block.type == "tool_use":
            return block.input
    raise RuntimeError(f"No structured extraction returned for {filename}")


# Title patterns for docs that are near-certainly not food content at all —
# kept deliberately narrow (a specific document category, not a vague
# keyword) to minimize the risk of dropping a real menu. Notably does NOT
# include anything like "invoice" or a plain client/pickup name: real
# per-event menus have turned up under exactly those filenames in this
# archive, so category alone can't be trusted to exclude them — only
# `extract`'s actual per-doc read of the content can.
CURATE_EXCLUDE_PATTERNS = [
    (r"credit.?card.?authoriz", "payment authorization form"),
    (r"\bw[-\s]?9\b", "tax form (W-9)"),
    (r"certificate.?of.?insurance|\bcoi\b|insurance.?cert", "insurance certificate"),
    (r"\bcontact.?list\b", "contact list"),
    (r"\bname.?tags?\b|\bplace.?cards?\b|\bnametags?\b", "name tag / place card template"),
    (r"\blogo(s)?\b", "logo file"),
    (r"^new microsoft word document", "blank/default Word filename"),
    (r"cutlery|glassware|linen[s]?\b.*(rental|hire)|china\s*,\s*tables", "rental equipment list"),
    (r"\bbusiness\s+card", "business card"),
    (r"\bflyer\b|\badvertisement\b|\bmarketing\b", "marketing collateral"),
    (r"\bfloor\s*plan\b|\bseating\s*chart\b", "floor plan / seating chart"),
]


def _curate_title_exclude_reason(name: str):
    low = name.lower()
    for pattern, reason in CURATE_EXCLUDE_PATTERNS:
        if re.search(pattern, low):
            return reason
    return None


_PREP_LIST_QTY_WORDS = re.compile(r"\b(pax|hotel\s*pan|hotel\s*paan|tray|case|bucket|packet|pouch)\b", re.I)
_PREP_LIST_DIVIDER = re.compile(r"\*{3,}")


def _looks_like_prep_list(text: str) -> bool:
    has_dollar = "$" in text
    qty_hits = len(_PREP_LIST_QTY_WORDS.findall(text))
    has_divider_structure = bool(_PREP_LIST_DIVIDER.search(text))
    return (not has_dollar) and qty_hits >= 2 and (not has_divider_structure)


def cmd_curate(args):
    """Turn a raw source archive into a curated folder safe to hand to
    `extract` — strips true junk, collapses exact-duplicate copies, and
    drops only the document categories that could never contain a menu.
    Never modifies or deletes anything under --input; --output is built
    fresh each run (deterministic, so re-running after --input changes is
    always safe)."""
    import hashlib
    import shutil
    from collections import defaultdict

    in_dir = Path(args.input)
    out_dir = Path(args.output)

    all_files = sorted(p for p in in_dir.rglob("*") if p.suffix.lower() in (".docx", ".pdf"))
    print(f"Found {len(all_files)} documents under {in_dir}")

    # 1. Junk: 0-byte files and Word's own "~$..." lock/temp files (created
    # while a doc is open in Word) — neither has any real content.
    junk = [p for p in all_files if p.name.startswith("~$") or p.stat().st_size == 0]
    junk_set = set(junk)
    real_files = [p for p in all_files if p not in junk_set]
    print(f"Junk (lock files / 0-byte): {len(junk)}")

    # 2. Exact-duplicate content, by hash — a real source archive commonly
    # has the same doc saved under multiple names/folders/years. Keep one
    # per identical-content group, preferring the copy whose name has no
    # "(2)"-style suffix, else the shortest path.
    def file_hash(p: Path) -> str:
        h = hashlib.sha1()
        with open(p, "rb") as f:
            for chunk in iter(lambda: f.read(1 << 20), b""):
                h.update(chunk)
        return h.hexdigest()

    by_hash = defaultdict(list)
    for p in real_files:
        by_hash[file_hash(p)].append(p)

    def pick_keeper(paths):
        no_paren = [p for p in paths if "(" not in p.stem]
        pool = no_paren if no_paren else paths
        return min(pool, key=lambda p: len(str(p)))

    deduped = []
    dropped_dupes = []
    for group in by_hash.values():
        if len(group) == 1:
            deduped.append(group[0])
        else:
            keep = pick_keeper(group)
            deduped.append(keep)
            dropped_dupes.extend(p for p in group if p != keep)
    print(f"Exact-duplicate copies dropped: {len(dropped_dupes)}")

    # 3. Title-based exclusion — only the categories in
    # CURATE_EXCLUDE_PATTERNS, deliberately conservative (see its comment).
    after_title = []
    title_flagged = []
    for p in deduped:
        reason = _curate_title_exclude_reason(p.name)
        if reason:
            title_flagged.append((p, reason))
        else:
            after_title.append(p)
    print(f"Flagged as 'not menu-related' by title: {len(title_flagged)}")

    # 4. Content-based exclusion: internal kitchen prep/quantity notes
    # ("Onion bhji  300", truck equipment checklists) rather than a real
    # client-facing menu. Unlike the title check, this reads actual doc
    # text — no dollar sign anywhere (real priced menus in this archive
    # almost always show pricing) + 2+ portion/quantity words (pax, hotel
    # pan, tray, case, bucket, packet, pouch) + none of this company's own
    # real-menu formatting convention (repeated "***"-style dividers
    # between items — see EXTRACTION_PROMPT). That third condition exists
    # because the first two alone caught a real, well-formatted wedding-
    # celebration menu that used "hotel pan" for live-station quantities
    # and had no literal "$" in this copy — found by spot-checking a
    # sample before trusting the count, same as the title filter.
    kept = []
    prep_list_flagged = []
    for p in after_title:
        try:
            text = extract_text(p)
        except Exception:
            kept.append(p)  # let extract's own error handling deal with it
            continue
        if _looks_like_prep_list(text):
            prep_list_flagged.append(p)
        else:
            kept.append(p)
    print(f"Flagged as internal prep/quantity note (not a menu): {len(prep_list_flagged)}")
    print(f"Final curated set: {len(kept)}")

    # Build --output fresh every run rather than incrementally patching it,
    # so it can never drift from what --input + these rules actually say.
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True)
    for p in kept:
        rel = p.relative_to(in_dir)
        dest = out_dir / rel
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(p, dest)

    report = {
        "input": str(in_dir),
        "output": str(out_dir),
        "total_files": len(all_files),
        "junk_count": len(junk),
        "junk_files": sorted(str(p.relative_to(in_dir)) for p in junk),
        "duplicate_dropped_count": len(dropped_dupes),
        "duplicate_dropped_files": sorted(str(p.relative_to(in_dir)) for p in dropped_dupes),
        "title_flagged_count": len(title_flagged),
        "title_flagged_files": sorted(
            [{"path": str(p.relative_to(in_dir)), "reason": r} for p, r in title_flagged],
            key=lambda d: (d["reason"], d["path"]),
        ),
        "prep_list_flagged_count": len(prep_list_flagged),
        "prep_list_flagged_files": sorted(str(p.relative_to(in_dir)) for p in prep_list_flagged),
        "kept_count": len(kept),
    }
    report_path = out_dir.parent / "curation_report.json"
    report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(f"\nCuration report written to: {report_path}")
    print(f"Curated set ready at: {out_dir}")


def _extract_one(client, in_dir: Path, out_dir: Path, path: Path) -> str:
    """Extract a single doc; returns a short status line. Never raises —
    every failure mode (bad file, API error, empty text) is caught here so
    one bad doc can't kill the rest of a concurrent batch."""
    import hashlib
    rel = path.relative_to(in_dir)
    # Two different source subfolders can each contain a file with the same
    # stem (e.g. "Menu 1.docx" under both a 2016/ and a 2018/ folder) —
    # stem-only output naming would collide and overwrite one's extraction.
    # Prefixing with a short hash of the full relative path keeps every
    # doc's output distinct while staying resumable (same input path always
    # hashes to the same output file).
    digest = hashlib.sha1(str(rel).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{digest}.json"
    if out_path.exists():
        return f"skip (already extracted): {rel}"
    try:
        text = extract_text(path)
        if not text.strip():
            return f"empty text, skipping: {rel}"
        result = call_claude_extract(client, path.name, text)
        # Full relative path, not just the filename — two different source
        # subfolders can share a filename, and this is what ends up in each
        # dish's traceable source_docs list later.
        result["_source_doc"] = str(rel)
        # encoding="utf-8" required — Path.write_text() otherwise falls back
        # to the system locale codepage (cp1252 on Windows), same root
        # cause as the dish_review.csv write below.
        out_path.write_text(json.dumps(result, indent=2), encoding="utf-8")
        return f"done: {rel}"
    except Exception as e:
        return f"FAILED: {rel} -> {e}"


def cmd_extract(args):
    import anthropic
    from concurrent.futures import ThreadPoolExecutor, as_completed

    # One client shared across worker threads — the SDK's underlying HTTP
    # client is safe for concurrent use (each call is an independent
    # request), so this doesn't need a client-per-thread.
    client = anthropic.Anthropic()

    in_dir = Path(args.input)
    out_dir = Path(args.output)
    out_dir.mkdir(parents=True, exist_ok=True)

    # rglob, not iterdir — real source archives (e.g. a Google Drive export)
    # are organized in subfolders (by year, in this case), not dumped flat
    # into one directory. iterdir() alone would silently find zero files
    # pointed at a folder like that.
    files = sorted([p for p in in_dir.rglob("*") if p.suffix.lower() in (".docx", ".pdf")])
    print(f"Found {len(files)} documents under {in_dir}")

    # This is I/O-bound (waiting on the Claude API over the network), not
    # CPU-bound, so a thread pool works fine despite the GIL — each file's
    # extraction is fully independent (own API call, own uniquely-named
    # output file), nothing shared to race on. Concurrency is deliberately
    # modest by default to stay under typical per-account rate limits;
    # raise --concurrency if your tier allows more.
    done = 0
    with ThreadPoolExecutor(max_workers=args.concurrency) as pool:
        futures = {pool.submit(_extract_one, client, in_dir, out_dir, p): p for p in files}
        for future in as_completed(futures):
            done += 1
            status = future.result()
            print(f"[{done}/{len(files)}] {status}")
        time.sleep(0.3)  # light rate-limit courtesy

    print(f"Done. Per-doc extractions in {out_dir}/")


def normalize_name(name: str) -> str:
    return "".join(c.lower() for c in name if c.isalnum() or c.isspace()).strip()


def cmd_aggregate(args):
    from rapidfuzz import fuzz

    in_dir = Path(args.input)
    out_dir = Path(args.output)
    out_dir.mkdir(parents=True, exist_ok=True)

    all_items = []  # flat list, one row per (item, source menu)
    for jf in sorted(in_dir.glob("*.json")):
        # encoding="utf-8" required to match cmd_extract's write — same
        # missing-default issue as the dish_review.csv write below.
        data = json.loads(jf.read_text(encoding="utf-8"))
        source_doc = data.get("_source_doc", jf.stem)
        for menu in data.get("menus", []):
            for item in menu.get("items", []):
                all_items.append({
                    **item,
                    "source_doc": source_doc,
                    "source_menu": menu.get("menu_label"),
                    "source_occasion_guess": menu.get("occasion_type_guess"),
                })

    print(f"Loaded {len(all_items)} raw item mentions from {len(list(in_dir.glob('*.json')))} docs")

    # Fuzzy-dedupe by normalized name, merging tags/sources, keeping the
    # highest-confidence veg/nonveg + course call across duplicates.
    merged = []
    used = [False] * len(all_items)
    CONF_RANK = {"high": 2, "medium": 1, "low": 0}

    for i, item in enumerate(all_items):
        if used[i]:
            continue
        group = [item]
        used[i] = True
        norm_i = normalize_name(item["name"])
        for j in range(i + 1, len(all_items)):
            if used[j]:
                continue
            norm_j = normalize_name(all_items[j]["name"])
            # token_sort_ratio alone misses "Masala Chai" vs "Freshly Brewed
            # Masala Chai" — it scores full sorted-token strings against each
            # other, so extra descriptive words drag the score down even
            # though one name is just a wordier version of the other.
            # token_set_ratio is built for exactly that containment case
            # (one name's words are a subset of the other's), so take the
            # best of both instead of relying on token_sort_ratio alone.
            similarity = max(
                fuzz.token_sort_ratio(norm_i, norm_j),
                fuzz.token_set_ratio(norm_i, norm_j),
            )
            if similarity >= 90:
                group.append(all_items[j])
                used[j] = True

        best = max(group, key=lambda x: CONF_RANK.get(x.get("confidence", "low"), 0))
        cuisine_tags = sorted({t for g in group for t in g.get("cuisine_tags_guess", [])})
        allergens = sorted({a for g in group for a in g.get("allergens_guess", [])})
        dietary_flags = sorted({d for g in group for d in g.get("dietary_flags_guess", [])})
        # Religion suitability: intersect across sightings rather than union —
        # if one sighting flags "not suitable for Hindu weddings" that has to
        # win over another sighting that just didn't think to check.
        religion_sets = [set(g.get("religion_suitability_guess", ["any"])) for g in group]
        religion_suitability = sorted(set.intersection(*religion_sets)) if religion_sets else ["any"]
        # Specific tags win over "any" whenever any sighting gave one — "any"
        # only means "this particular sighting had no specific signal," not
        # "broadly appropriate despite what other sightings say." A plain
        # union let one generic-menu sighting's "any" outlive every other
        # sighting's specific tags forever (found by generating real menus:
        # dinner curries kept showing up for breakfast because one source
        # doc's guess of "any" never got displaced by the 3+ other sightings
        # that correctly said dinner_reception/lunch/sangeet/wedding_lunch).
        # "any" only survives if literally every sighting ever guessed it.
        occasion_sets = [set(g.get("occasion_suitability_guess", [])) for g in group]
        specific_occasion_tags = {o for s in occasion_sets for o in s if o != "any"}
        occasion_suitability = sorted(specific_occasion_tags) if specific_occasion_tags else ["any"]
        sources = sorted({g["source_doc"] for g in group})

        merged.append({
            "name": best["name"],
            "course": best["course"],
            "veg_nonveg": best["veg_nonveg_guess"],
            "cuisine_tags": cuisine_tags,
            "price_weight": best["price_weight_guess"],
            "is_staple": any(g.get("is_staple_guess") for g in group),
            "allergens": allergens,
            "dietary_flags": dietary_flags,
            "religion_suitability": religion_suitability or ["any"],
            "occasion_suitability": occasion_suitability or ["any"],
            "spice_level": best.get("spice_level_guess"),
            "prep_method": best.get("prep_method_guess"),
            "tax_category": best.get("tax_category_guess", "prepared_food"),
            "confidence": best["confidence"],
            "seen_in_docs_count": len(sources),
            "example_sources": sources[:3],
        })

    merged.sort(key=lambda x: (x["confidence"] != "low", x["confidence"] != "medium", -x["seen_in_docs_count"]))

    # Review CSV — human-eyeball pass before this becomes the live database.
    import csv
    review_path = out_dir / "dish_review.csv"
    # encoding="utf-8" is required, not a default — open() with no encoding
    # falls back to the system locale codepage (cp1252 on this Windows
    # machine), not UTF-8, despite _read_reviewed_csv's comment claiming
    # this file is "always written as UTF-8." Found by generating real
    # menus: accented dish names (sautéed, etc.) were coming out mangled
    # even in a file nobody had opened in Excel yet — the corruption was
    # happening right here on write, before any human ever touched it.
    with open(review_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["name", "course", "veg_nonveg", "cuisine_tags", "price_weight", "is_staple",
                    "allergens", "dietary_flags", "religion_suitability", "occasion_suitability",
                    "spice_level", "prep_method", "tax_category",
                    "confidence", "seen_in_docs_count", "example_sources"])
        for m in merged:
            w.writerow([m["name"], m["course"], m["veg_nonveg"], "|".join(m["cuisine_tags"]), m["price_weight"], m["is_staple"],
                        "|".join(m["allergens"]), "|".join(m["dietary_flags"]), "|".join(m["religion_suitability"]), "|".join(m["occasion_suitability"]),
                        m.get("spice_level") or "", m.get("prep_method") or "", m["tax_category"],
                        m["confidence"], m["seen_in_docs_count"], "; ".join(m["example_sources"])])

    seed_path = out_dir / "menu_items_seed.json"
    seed_path.write_text(json.dumps(merged, indent=2), encoding="utf-8")

    low_conf = sum(1 for m in merged if m["confidence"] == "low")
    print(f"Merged into {len(merged)} unique dishes ({low_conf} flagged low-confidence).")
    print(f"Review spreadsheet: {review_path}")
    print(f"Seed data for the app: {seed_path}")


DEFAULT_TAX_CATEGORIES = [
    # name, jurisdiction, rate_percent — placeholder rates, replace with your
    # actual state/county figures before this feeds real quotes.
    ("prepared_food", "FL", 7.000),
    ("alcohol", "FL", 7.000),
    ("equipment_rental", "FL", 7.000),
    ("service_labor", "FL", 0.000),
]


def _read_reviewed_csv(csv_path: Path) -> list:
    """Read dish_review.csv — the file a human actually opens and edits —
    back into the same shape cmd_load expects. A row the reviewer deleted
    (e.g. a duplicate fuzzy matching couldn't catch, like two totally
    differently-worded names for the same dish) simply won't appear here,
    which is the point: the reviewed CSV, not the raw un-reviewed seed JSON,
    is what actually determines what gets loaded."""
    import csv

    def split_pipe(value: str) -> list:
        return [v for v in value.split("|") if v]

    # This script always *writes* dish_review.csv as UTF-8, but the reviewer
    # opens and edits it in Excel — which, on Windows, saves "CSV (Comma
    # delimited)" back out using the system's ANSI code page (cp1252 for
    # US/Western-Europe locales), not UTF-8. Any accented character in a
    # dish name (sautéed, café, jalapeño...) silently changes encoding on
    # that save, which then fails a strict UTF-8 read. Try UTF-8 first
    # (covers an unedited file, and utf-8-sig handles the BOM Excel adds),
    # fall back to cp1252 — the encoding Excel actually used — instead of
    # erroring out on a normal review-and-save round trip.
    raw_rows = None
    last_error = None
    for encoding in ("utf-8-sig", "cp1252"):
        try:
            with open(csv_path, newline="", encoding=encoding) as f:
                raw_rows = list(csv.DictReader(f))
            break
        except UnicodeDecodeError as e:
            last_error = e
    if raw_rows is None:
        raise last_error

    items = []
    for row in raw_rows:
        items.append({
            "name": row["name"],
            "course": row["course"],
            "veg_nonveg": row["veg_nonveg"],
            "cuisine_tags": split_pipe(row["cuisine_tags"]),
            "price_weight": row["price_weight"],
            "is_staple": row["is_staple"].strip().lower() == "true",
            "allergens": split_pipe(row["allergens"]),
            "dietary_flags": split_pipe(row["dietary_flags"]),
            "religion_suitability": split_pipe(row["religion_suitability"]) or ["any"],
            "occasion_suitability": split_pipe(row["occasion_suitability"]) or ["any"],
            "spice_level": row["spice_level"] or None,
            "prep_method": row["prep_method"] or None,
            "tax_category": row["tax_category"] or None,
            "confidence": row["confidence"] or None,
            "example_sources": [s for s in row["example_sources"].split("; ") if s],
        })
    return items


def _strip_prisma_only_params(database_url: str) -> str:
    """DATABASE_URL as used by the Next.js app carries query params that
    are Prisma-engine conventions, not real libpq/Postgres options —
    `pgbouncer=true` (tells Prisma to skip prepared statements, since the
    Supabase pooler runs in transaction mode) is the one that's bitten us;
    `connection_limit`/`pool_timeout`/`schema` are the other Prisma-only
    ones worth guarding against. psycopg2's DSN parser rejects any query
    param it doesn't recognize, so this script — which connects with
    plain psycopg2, not Prisma — needs them stripped before connecting."""
    from urllib.parse import urlsplit, urlunsplit, parse_qsl, urlencode

    prisma_only = {"pgbouncer", "connection_limit", "pool_timeout", "schema"}
    parts = urlsplit(database_url)
    kept = [(k, v) for k, v in parse_qsl(parts.query) if k not in prisma_only]
    return urlunsplit((parts.scheme, parts.netloc, parts.path, urlencode(kept), parts.fragment))


def _ensure_tax_categories(cur) -> dict:
    """Upsert the default tax categories and return {name: id} — every
    upsert path (load/stage/promote) needs this same lookup, since dishes
    reference a tax category by name, not id, until this resolves it."""
    for name, jurisdiction, rate in DEFAULT_TAX_CATEGORIES:
        cur.execute(
            """
            INSERT INTO tax_categories (id, name, jurisdiction, rate_percent, effective_date)
            VALUES (gen_random_uuid(), %s, %s, %s, now())
            ON CONFLICT (name) DO NOTHING
            """,
            (name, jurisdiction, rate),
        )
    cur.execute("SELECT id, name FROM tax_categories")
    return {name: id_ for id_, name in cur.fetchall()}


def _upsert_menu_item(cur, item: dict, tax_id_by_name: dict) -> None:
    """Insert-or-update one dish into the live `menu_items` table, upserted
    by its unique name. Shared by `load` (source: a human-reviewed CSV) and
    `promote` (source: approved/edited MenuItemDraft rows) — both hand this
    the same dict shape, so the actual SQL only needs to exist once."""
    tax_category_id = tax_id_by_name.get(item.get("tax_category", "prepared_food"))
    cur.execute(
        """
        INSERT INTO menu_items (
            id, name, course, veg_nonveg, cuisine_tags, price_weight, is_staple,
            allergens, dietary_flags, religion_suitability, occasion_suitability,
            spice_level, prep_method, tax_category_id, confidence, source_docs,
            active, created_at, updated_at
        ) VALUES (
            gen_random_uuid(), %(name)s, %(course)s, %(veg_nonveg)s, %(cuisine_tags)s, %(price_weight)s, %(is_staple)s,
            %(allergens)s, %(dietary_flags)s, %(religion_suitability)s, %(occasion_suitability)s,
            %(spice_level)s, %(prep_method)s, %(tax_category_id)s, %(confidence)s, %(source_docs)s,
            true, now(), now()
        )
        ON CONFLICT (name) DO UPDATE SET
            course = EXCLUDED.course,
            veg_nonveg = EXCLUDED.veg_nonveg,
            cuisine_tags = EXCLUDED.cuisine_tags,
            price_weight = EXCLUDED.price_weight,
            is_staple = EXCLUDED.is_staple,
            allergens = EXCLUDED.allergens,
            dietary_flags = EXCLUDED.dietary_flags,
            religion_suitability = EXCLUDED.religion_suitability,
            occasion_suitability = EXCLUDED.occasion_suitability,
            spice_level = EXCLUDED.spice_level,
            prep_method = EXCLUDED.prep_method,
            tax_category_id = EXCLUDED.tax_category_id,
            confidence = EXCLUDED.confidence,
            source_docs = EXCLUDED.source_docs,
            updated_at = now()
        """,
        {
            "name": item["name"],
            "course": item["course"],
            "veg_nonveg": item["veg_nonveg"],
            "cuisine_tags": item["cuisine_tags"],
            "price_weight": item["price_weight"],
            "is_staple": item["is_staple"],
            "allergens": item.get("allergens", []),
            "dietary_flags": item.get("dietary_flags", []),
            "religion_suitability": item.get("religion_suitability", ["any"]),
            "occasion_suitability": item.get("occasion_suitability", ["any"]),
            "spice_level": item.get("spice_level"),
            "prep_method": item.get("prep_method"),
            "tax_category_id": tax_category_id,
            "confidence": item["confidence"],
            "source_docs": item.get("example_sources", []),
        },
    )


def cmd_load(args):
    import psycopg2

    review_path = Path(args.input) / "dish_review.csv"
    items = _read_reviewed_csv(review_path)
    print(f"Loading {len(items)} dishes from {review_path}")

    conn = psycopg2.connect(_strip_prisma_only_params(args.database_url))
    conn.autocommit = False
    try:
        with conn.cursor() as cur:
            tax_id_by_name = _ensure_tax_categories(cur)
            for item in items:
                _upsert_menu_item(cur, item, tax_id_by_name)
        conn.commit()
        print(f"Loaded {len(items)} dishes into menu_items (upserted by unique name).")
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def cmd_stage(args):
    """Push aggregate's output (menu_items_seed.json) into menu_item_drafts
    — every dish starts life here as review_status='pending', never landing
    in the live menu_items table until a human approves it through the
    /review web UI and `promote` runs. Safe to re-run after a corrected
    re-extraction: upserted by name, same as load/promote."""
    import psycopg2

    seed_path = Path(args.input) / "menu_items_seed.json"
    items = json.loads(seed_path.read_text(encoding="utf-8"))
    print(f"Staging {len(items)} dishes from {seed_path}")

    conn = psycopg2.connect(_strip_prisma_only_params(args.database_url))
    conn.autocommit = False
    try:
        with conn.cursor() as cur:
            tax_id_by_name = _ensure_tax_categories(cur)
            for item in items:
                tax_category_id = tax_id_by_name.get(item.get("tax_category", "prepared_food"))
                cur.execute(
                    """
                    INSERT INTO menu_item_drafts (
                        id, name, course, veg_nonveg, cuisine_tags, price_weight, is_staple,
                        allergens, dietary_flags, religion_suitability, occasion_suitability,
                        spice_level, prep_method, tax_category_id, confidence, source_docs,
                        review_status, created_at, updated_at
                    ) VALUES (
                        gen_random_uuid(), %(name)s, %(course)s, %(veg_nonveg)s, %(cuisine_tags)s, %(price_weight)s, %(is_staple)s,
                        %(allergens)s, %(dietary_flags)s, %(religion_suitability)s, %(occasion_suitability)s,
                        %(spice_level)s, %(prep_method)s, %(tax_category_id)s, %(confidence)s, %(source_docs)s,
                        'pending', now(), now()
                    )
                    -- Re-staging (e.g. a corrected re-extraction) resets the review back to
                    -- pending rather than silently keeping a stale prior decision around —
                    -- the content changed, so any earlier approve/reject no longer applies.
                    ON CONFLICT (name) DO UPDATE SET
                        course = EXCLUDED.course,
                        veg_nonveg = EXCLUDED.veg_nonveg,
                        cuisine_tags = EXCLUDED.cuisine_tags,
                        price_weight = EXCLUDED.price_weight,
                        is_staple = EXCLUDED.is_staple,
                        allergens = EXCLUDED.allergens,
                        dietary_flags = EXCLUDED.dietary_flags,
                        religion_suitability = EXCLUDED.religion_suitability,
                        occasion_suitability = EXCLUDED.occasion_suitability,
                        spice_level = EXCLUDED.spice_level,
                        prep_method = EXCLUDED.prep_method,
                        tax_category_id = EXCLUDED.tax_category_id,
                        confidence = EXCLUDED.confidence,
                        source_docs = EXCLUDED.source_docs,
                        review_status = 'pending',
                        reviewed_at = NULL,
                        updated_at = now()
                    """,
                    {
                        "name": item["name"],
                        "course": item["course"],
                        "veg_nonveg": item["veg_nonveg"],
                        "cuisine_tags": item["cuisine_tags"],
                        "price_weight": item["price_weight"],
                        "is_staple": item["is_staple"],
                        "allergens": item.get("allergens", []),
                        "dietary_flags": item.get("dietary_flags", []),
                        "religion_suitability": item.get("religion_suitability", ["any"]),
                        "occasion_suitability": item.get("occasion_suitability", ["any"]),
                        "spice_level": item.get("spice_level"),
                        "prep_method": item.get("prep_method"),
                        "tax_category_id": tax_category_id,
                        "confidence": item.get("confidence"),
                        "source_docs": item.get("example_sources", []),
                    },
                )
        conn.commit()
        print(f"Staged {len(items)} dishes into menu_item_drafts (review_status='pending').")
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def cmd_promote(args):
    """Take whatever the customer approved/edited at /review and upsert it
    into the live menu_items table — the "tomorrow" batch step. Anything
    still 'pending' or 'rejected' is left alone; nothing not explicitly
    reviewed ever reaches the live menu-generation wizard."""
    import psycopg2
    import psycopg2.extras

    conn = psycopg2.connect(_strip_prisma_only_params(args.database_url))
    conn.autocommit = False
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            tax_id_by_name = _ensure_tax_categories(cur)
            # Need the tax category's *name*, not its id, since _upsert_menu_item
            # takes item["tax_category"] as a name and re-resolves the id itself.
            id_to_tax_name = {id_: name for name, id_ in tax_id_by_name.items()}

            cur.execute(
                """
                SELECT id, name, course, veg_nonveg, cuisine_tags, price_weight, is_staple,
                       allergens, dietary_flags, religion_suitability, occasion_suitability,
                       spice_level, prep_method, tax_category_id, confidence, source_docs
                FROM menu_item_drafts
                WHERE review_status IN ('approved', 'edited')
                """
            )
            drafts = cur.fetchall()
            print(f"Promoting {len(drafts)} approved/edited drafts to menu_items")

            promoted_ids = []
            for draft in drafts:
                item = {
                    "name": draft["name"],
                    "course": draft["course"],
                    "veg_nonveg": draft["veg_nonveg"],
                    "cuisine_tags": draft["cuisine_tags"] or [],
                    "price_weight": draft["price_weight"],
                    "is_staple": draft["is_staple"],
                    "allergens": draft["allergens"] or [],
                    "dietary_flags": draft["dietary_flags"] or [],
                    "religion_suitability": draft["religion_suitability"] or ["any"],
                    "occasion_suitability": draft["occasion_suitability"] or ["any"],
                    "spice_level": draft["spice_level"],
                    "prep_method": draft["prep_method"],
                    "tax_category": id_to_tax_name.get(draft["tax_category_id"], "prepared_food"),
                    "confidence": draft["confidence"],
                    "example_sources": draft["source_docs"] or [],
                }
                _upsert_menu_item(cur, item, tax_id_by_name)
                promoted_ids.append(draft["id"])

            if promoted_ids:
                # Kept, not deleted — an audit trail of what this batch actually
                # contained and what a human decided about it (see MenuItemDraft's
                # schema.prisma comment).
                cur.execute(
                    "UPDATE menu_item_drafts SET review_status = 'promoted', updated_at = now() WHERE id = ANY(%s)",
                    (promoted_ids,),
                )
        conn.commit()
        print(f"Promoted {len(drafts)} dishes into menu_items; drafts marked 'promoted'.")
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = parser.add_subparsers(dest="command", required=True)

    p_curate = sub.add_parser("curate", help="Strip junk/duplicates/non-menu docs from a raw source archive before extract.")
    p_curate.add_argument("--input", required=True, help="Folder of raw .docx/.pdf source documents (not modified).")
    p_curate.add_argument("--output", required=True, help="Folder to write the curated copy into (rebuilt fresh each run).")
    p_curate.set_defaults(func=cmd_curate)

    p_extract = sub.add_parser("extract", help="Run per-document LLM extraction (resumable, skips docs already done).")
    p_extract.add_argument("--input", required=True, help="Folder of .docx/.pdf menu documents.")
    p_extract.add_argument("--output", required=True, help="Folder to write one JSON file per source doc.")
    p_extract.add_argument("--concurrency", type=int, default=5,
                            help="How many docs to extract in parallel (default 5). Each is one Claude API "
                                 "call over the network, so this is I/O-bound — raise it if your account's "
                                 "rate limit allows more throughput.")
    p_extract.set_defaults(func=cmd_extract)

    p_agg = sub.add_parser("aggregate", help="Merge, dedupe, and produce the reviewable seed dataset.")
    p_agg.add_argument("--input", required=True, help="Folder of per-doc JSON from the extract step.")
    p_agg.add_argument("--output", required=True, help="Folder to write dish_review.csv and menu_items_seed.json.")
    p_agg.set_defaults(func=cmd_aggregate)

    p_load = sub.add_parser("load", help="Load the seed dataset into Postgres (tables from prisma/schema.prisma).")
    p_load.add_argument("--input", required=True, help="Folder containing dish_review.csv (the aggregate step's output, after your human review pass).")
    p_load.add_argument("--database-url", required=True, help="Postgres connection string, e.g. $DATABASE_URL.")
    p_load.set_defaults(func=cmd_load)

    p_stage = sub.add_parser("stage", help="Push aggregate's output into menu_item_drafts for web review at /review (pending, not live yet).")
    p_stage.add_argument("--input", required=True, help="Folder containing menu_items_seed.json (the aggregate step's output).")
    p_stage.add_argument("--database-url", required=True, help="Postgres connection string, e.g. $DATABASE_URL.")
    p_stage.set_defaults(func=cmd_stage)

    p_promote = sub.add_parser("promote", help="Upsert every approved/edited menu_item_drafts row into the live menu_items table.")
    p_promote.add_argument("--database-url", required=True, help="Postgres connection string, e.g. $DATABASE_URL.")
    p_promote.set_defaults(func=cmd_promote)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
